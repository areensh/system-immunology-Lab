from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    return h

def add_body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_bold_body(bold_part, rest):
    p = doc.add_paragraph()
    run = p.add_run(bold_part)
    run.bold = True
    p.add_run(rest)
    p.paragraph_format.space_after = Pt(6)

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Shading Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()

def add_caption(text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

# ============================================================
# TITLE PAGE
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading('Immune Repertoire Analysis', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_heading('Results', level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in subtitle.runs:
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('\nDate: ').bold = True
info.add_run('June 15, 2026\n')
info.add_run('Data Source: ').bold = True
info.add_run('iReceptor Statistics API\n')
info.add_run('Total Repertoires: ').bold = True
info.add_run('100 (85 with complete metadata)')

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS
# ============================================================
add_heading('Table of Contents', level=1)
toc_items = [
    "1. Dataset Discovery",
    "2. Metadata Exploration",
    "   2.1 Available Metadata Categories",
    "   2.2 Study Metadata",
    "   2.3 Subject Metadata",
    "   2.4 Diagnosis and Intervention Metadata",
    "   2.5 Sample Metadata",
    "   2.6 Process Metadata",
    "3. Cohort Characterization",
    "   3.1 Disease Category Distribution",
    "   3.2 Sex Distribution",
    "   3.3 Age Distribution",
    "   3.4 Cross-Stratifications",
    "4. Statistical Endpoints",
    "5. Clonal Diversity and Expansion",
    "   5.1 Clone Count",
    "   5.2 Clone Size",
    "6. Repertoire Dominance",
    "   6.1 Top X Clone Fraction",
    "7. CDR3 Structural Features",
    "   7.1 CDR3 Length",
    "8. Summary of Key Findings",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)

doc.add_page_break()

# ============================================================
# SECTION 1: DATASET DISCOVERY
# ============================================================
add_heading('1. Dataset Discovery', level=1)

add_body(
    'The iReceptor Statistics API provides a discovery-driven interface for exploring '
    'immune repertoire data. As a first step, we queried the API to retrieve all available '
    'datasets in the system.'
)

add_body(
    'The query returned 5 available datasets, of which 3 are related to COVID-19 studies. '
    'We selected these COVID-19 datasets as a use case because repertoire studies in infectious '
    'disease are a key application of AIRR-seq, and COVID-19 provides a well-characterized '
    'clinical spectrum (naive, mild, severe, recovered) that allows us to demonstrate the '
    'analytical capabilities of the platform across biologically meaningful comparisons.'
)

add_table(
    ['Dataset', 'Study Title', 'Relevant Publication'],
    [
        ['covid_vaccine', 'covid_vaccine', 'PMID: 33858945'],
        ['covid19-2', 'covid19-2', 'PMID: 33384691'],
        ['Covid19', 'Covid19', 'PMID: 32669287'],
        ['covid_vaccine2', 'covid_vaccine2', 'PMID: 34648302'],
        ['patterns in COVID-19', 'patterns in COVID-19', 'PRJNA839749'],
    ]
)

add_caption('Figure 1.1. Number of subjects per dataset. [See figure file]')

doc.add_page_break()

# ============================================================
# SECTION 2: METADATA EXPLORATION
# ============================================================
add_heading('2. Metadata Exploration', level=1)

add_body(
    'Before performing statistical analyses, we explored the metadata structure of the selected '
    'datasets using the API\'s metadata discovery endpoints. The /meta/categories/ endpoint '
    'returns the broad categories of metadata available, and /meta/categories/:type returns '
    'the specific fields and values within each category.'
)

# 2.1 Available categories
add_heading('2.1 Available Metadata Categories', level=2)

add_body('The API exposes 5 metadata categories:')

add_table(
    ['#', 'Category', 'Description'],
    [
        ['1', 'Study', 'Study-level information (title, publications, lab)'],
        ['2', 'Subject', 'Subject demographics (age, sex, ID)'],
        ['3', 'Diagnosis and Intervention', 'Disease stage, diagnosis, treatments'],
        ['4', 'Sample', 'Sample collection details'],
        ['5', 'Process', 'Sequencing and processing parameters'],
    ]
)

# 2.2 Study metadata
add_heading('2.2 Study Metadata', level=2)

add_body('Querying /meta/categories/Study returned the following fields and values:')

add_table(
    ['Field', 'Available Values'],
    [
        ['study_title', 'Covid19, covid_vaccine, covid19-2, patterns in COVID-19, covid_vaccine2'],
        ['Relevant publications', 'PMID: 32669287, PMID: 33858945, PMID: 33384691, PRJNA839749, PMID: 34648302'],
        ['study_ID', 'PRJNA638224'],
        ['lab_name', 'poria, Rabin, Shaare Zedek'],
    ]
)

# 2.3 Subject metadata
add_heading('2.3 Subject Metadata', level=2)
add_body('Querying /meta/categories/Subject returned the following fields and values:')
add_table(
    ['Field', 'Available Values'],
    [
        ['', ''],
        ['', ''],
        ['', ''],
    ]
)

# 2.4 Diagnosis
add_heading('2.4 Diagnosis and Intervention Metadata', level=2)
add_body('Querying /meta/categories/Diagnosis and Intervention returned the following fields and values:')
add_table(
    ['Field', 'Available Values'],
    [
        ['', ''],
        ['', ''],
    ]
)

# 2.5 Sample
add_heading('2.5 Sample Metadata', level=2)
add_body('Querying /meta/categories/Sample returned the following fields and values:')
add_table(
    ['Field', 'Available Values'],
    [
        ['', ''],
        ['', ''],
    ]
)

# 2.6 Process
add_heading('2.6 Process Metadata', level=2)
add_body('Querying /meta/categories/process returned the following fields and values:')
add_table(
    ['Field', 'Available Values'],
    [
        ['', ''],
        ['', ''],
    ]
)

doc.add_page_break()

# ============================================================
# SECTION 3: COHORT CHARACTERIZATION
# ============================================================
add_heading('3. Cohort Characterization', level=1)

add_body(
    'Using the metadata endpoint, we queried subject-level information stratified by disease stage, '
    'sex, and age across the 3 selected COVID-19 datasets. Of the 100 repertoires retrieved, '
    '85 had complete metadata and were included in all downstream analyses.'
)

# 3.1 Disease
add_heading('3.1 Disease Category Distribution', level=2)

add_body(
    'The raw disease_stage field contained 13 distinct annotations across studies. '
    'These were harmonized into 5 biologically meaningful categories:'
)

add_table(
    ['Raw Disease Stage(s)', 'Harmonized Category', 'N'],
    [
        ['COVID Naive', 'COVID Naive', '8'],
        ['mild, non-severe', 'Mild', '38'],
        ['Early phase-Stable, Early phase-Improving, Recovering...', 'Moderate', '9'],
        ['Recovered, COVID recovered, Recovering post-ICU', 'Recovered', '12'],
        ['severe, Early phase hypoxaemia', 'Severe', '18'],
    ]
)

add_body(
    'The Mild group is the largest (n=38, 45%), followed by Severe (n=18, 21%), '
    'Recovered (n=12, 14%), Moderate (n=9, 11%), and COVID Naive (n=8, 9%).'
)

add_caption('Figure 3.1. Number of subjects per disease category. [See figure file]')

# 3.2 Sex
add_heading('3.2 Sex Distribution', level=2)
add_body('The cohort has a slight male skew: 51 males (60%) and 34 females (40%).')
add_caption('Figure 3.2. Number of subjects by sex. [See figure file]')

# 3.3 Age
add_heading('3.3 Age Distribution', level=2)
add_body('Age ranges from 18 to 88 years (median ~ 55). The cohort skews older, with 59% aged 51+.')

add_table(
    ['Age Group', 'N', '%'],
    [
        ['18-30', '10', '12%'],
        ['31-50', '25', '29%'],
        ['51-65', '29', '34%'],
        ['66+', '21', '25%'],
    ]
)
add_caption('Figure 3.3a. Age distribution histogram. [See figure file]')
add_caption('Figure 3.3b. Subjects per age group. [See figure file]')

# 3.4 Cross-strat
add_heading('3.4 Cross-Stratifications', level=2)
add_body(
    'Males are overrepresented in Severe (12M vs 6F) and Mild (22M vs 16F) groups. '
    'Older participants (51-65, 66+) are concentrated in the Severe group. '
    'Disease category confounds with study of origin, which is an important consideration.'
)
add_caption('Figure 3.4a. Disease category by sex. [See figure file]')
add_caption('Figure 3.4b. Disease category by age group. [See figure file]')
add_caption('Figure 3.4c. Study composition by disease category. [See figure file]')
add_caption('Figure 3.4d. Age distribution by disease category (boxplot). [See figure file]')
add_caption('Figure 3.4e. Heatmap: disease x age x sex. [See figure file]')

doc.add_page_break()

# ============================================================
# SECTION 4: STATISTICAL ENDPOINTS
# ============================================================
add_heading('4. Statistical Endpoints', level=1)

add_body(
    'Having characterized the cohort, we next queried the API to discover what statistical '
    'analyses are available. The /endpoints endpoint returned the following:'
)

add_table(
    ['#', 'Endpoint', 'Description', 'Metric'],
    [
        ['1', 'metaData', 'Subject metadata (age, sex, disease stage)', 'Counts and distributions'],
        ['2', 'clone_count', 'Number of unique clones per repertoire', 'Clonal diversity'],
        ['3', 'CDR3_length', 'Average CDR3 amino acid length of top clones', 'Structural features'],
        ['4', 'mutation', 'Somatic hypermutation statistics', 'Affinity maturation'],
    ]
)

add_body(
    'Each endpoint can be further explored using /endpoints/statistics/:type to retrieve '
    'the specific statistics available. For the clone_count endpoint, available statistics '
    'include clone count, clone size, and top X clone copy fractions.'
)

add_body(
    'The following sections present results from each statistical endpoint, organized by '
    'biological theme.'
)

doc.add_page_break()

# ============================================================
# SECTION 5: CLONAL DIVERSITY AND EXPANSION
# ============================================================
add_heading('5. Clonal Diversity and Expansion', level=1)

# 5.1 Clone Count
add_heading('5.1 Clone Count', level=2)

add_body(
    'Clone count represents the number of unique B-cell clones identified in each repertoire. '
    'Higher clone counts reflect greater clonal diversity.'
)

add_table(
    ['Disease Category', 'N', 'Median', 'Mean', 'Min', 'Max'],
    [
        ['COVID Naive', '8', '16,682', '18,086', '8,156', '29,178'],
        ['Mild', '38', '2,353', '2,593', '95', '5,997'],
        ['Moderate', '9', '5,085', '9,327', '3,459', '22,480'],
        ['Recovered', '12', '21,527', '19,661', '5,713', '33,699'],
        ['Severe', '18', '3,423', '6,782', '1,319', '22,452'],
    ]
)

add_bold_body('Key finding: ',
    'Recovered and COVID Naive individuals have the highest clone counts (median ~17K-22K), '
    'suggesting greater clonal diversity. Mild cases have the lowest (median ~2,350), likely '
    'reflecting differences in sequencing depth across studies.'
)

add_caption('Figure 5.1a. Clone count by disease category. [See figure file]')
add_caption('Figure 5.1b. Clone count by sex. [See figure file]')
add_caption('Figure 5.1c. Clone count by age group. [See figure file]')
add_caption('Figure 5.1d. Clone count by disease x sex. [See figure file]')
add_caption('Figure 5.1e. Clone count by disease x age. [See figure file]')

# 5.2 Clone Size
add_heading('5.2 Clone Size', level=2)

add_body(
    'Clone size represents the mean number of sequences per clone within each repertoire. '
    'Larger clone sizes indicate greater clonal expansion.'
)

add_table(
    ['Disease Category', 'N', 'Median (Mean CS)', 'Mean (Mean CS)'],
    [
        ['COVID Naive', '8', '114.0', '126.0'],
        ['Mild', '38', '36.4', '48.5'],
        ['Moderate', '9', '63.5', '60.0'],
        ['Recovered', '12', '66.9', '79.3'],
        ['Severe', '18', '53.4', '52.0'],
    ]
)

add_bold_body('Key finding: ',
    'COVID Naive individuals show the highest mean clone size (median 114), 3x higher than '
    'Mild cases (36.4). Severe and Moderate groups show intermediate expansion.'
)

add_caption('Figure 5.2a. Mean clone size by disease category. [See figure file]')
add_caption('Figure 5.2b. Mean clone size by sex. [See figure file]')
add_caption('Figure 5.2c. Mean clone size by age group. [See figure file]')
add_caption('Figure 5.2d. Mean clone size by disease x sex. [See figure file]')

doc.add_page_break()

# ============================================================
# SECTION 6: REPERTOIRE DOMINANCE
# ============================================================
add_heading('6. Repertoire Dominance', level=1)

add_heading('6.1 Top X Clone Fraction', level=2)

add_body(
    'The top X clone fraction measures the proportion of total sequence copies accounted for '
    'by the top 10, 100, or 1,000 clones. Higher fractions indicate greater oligoclonal '
    'dominance.'
)

add_table(
    ['Disease Category', 'N', 'Top 10 (median)', 'Top 100 (median)', 'Top 1000 (median)'],
    [
        ['COVID Naive', '8', '14.4%', '61.0%', '83.3%'],
        ['Mild', '38', '10.5%', '30.4%', '72.7%'],
        ['Moderate', '9', '22.6%', '45.5%', '74.1%'],
        ['Recovered', '12', '16.2%', '40.3%', '76.9%'],
        ['Severe', '18', '16.6%', '35.6%', '74.4%'],
    ]
)

add_bold_body('Key findings:', '')
bullets = [
    'Moderate disease shows the highest Top 10 dominance (22.6%), suggesting highly focused clonal responses.',
    'COVID Naive has the highest Top 100 and Top 1000 fractions (61% and 83%), indicating an oligoclonal baseline repertoire.',
    'Mild cases show the lowest dominance across all tiers (10.5% Top 10), consistent with a more polyclonal response.',
]
for b in bullets:
    p = doc.add_paragraph(b, style='List Bullet')
    p.paragraph_format.space_after = Pt(4)

add_caption('Figure 6.1a. Top X clone fractions by disease (faceted). [See figure file]')
add_caption('Figure 6.1b. Top X clone fractions by sex. [See figure file]')
add_caption('Figure 6.1c. Top X clone fractions by age group. [See figure file]')
add_caption('Figure 6.1d. Top 10 fraction by disease x sex. [See figure file]')
add_caption('Figure 6.1e. Top 10 fraction by disease x age. [See figure file]')

doc.add_page_break()

# ============================================================
# SECTION 7: CDR3 STRUCTURAL FEATURES
# ============================================================
add_heading('7. CDR3 Structural Features', level=1)

add_heading('7.1 CDR3 Length', level=2)

add_body(
    'CDR3 (Complementarity-Determining Region 3) length is a key structural feature of '
    'antibodies that determines antigen-binding specificity. Longer CDR3 regions can access '
    'recessed epitopes and are often associated with broadly neutralizing antibodies.'
)

add_table(
    ['Disease Category', 'N', 'Top 10 (median)', 'Top 100 (median)', 'Top 1000 (median)'],
    [
        ['COVID Naive', '8', '15.0 AA', '16.8 AA', '17.4 AA'],
        ['Mild', '38', '16.5 AA', '16.9 AA', '17.0 AA'],
        ['Moderate', '9', '17.3 AA', '17.2 AA', '17.5 AA'],
        ['Recovered', '12', '15.8 AA', '17.2 AA', '17.5 AA'],
        ['Severe', '18', '17.6 AA', '17.4 AA', '17.2 AA'],
    ]
)

add_bold_body('Key finding: ',
    'Severe cases show the longest CDR3 in their top 10 clones (median 17.6 AA), while '
    'COVID Naive show the shortest (15.0 AA). CDR3 lengths converge across groups at the '
    'Top 1000 level (~17.0-17.5 AA), indicating differences are driven by the most dominant clones.'
)

add_caption('Figure 7.1a. CDR3 length by disease (faceted by tier). [See figure file]')
add_caption('Figure 7.1b. CDR3 length by sex. [See figure file]')
add_caption('Figure 7.1c. CDR3 length by age group. [See figure file]')
add_caption('Figure 7.1d. CDR3 length (Top 10) by disease x sex. [See figure file]')
add_caption('Figure 7.1e. CDR3 length (Top 10) by disease x age. [See figure file]')
add_caption('Figure 7.1f. CDR3 length across clone tiers. [See figure file]')

doc.add_page_break()

# ============================================================
# SECTION 8: SUMMARY
# ============================================================
add_heading('8. Summary of Key Findings', level=1)

findings = [
    ('1. Cohort composition: ',
     '85 participants across 5 disease categories (Mild is the largest group at 45%). '
     'The cohort skews male (60%) and older (59% aged 51+). Severe cases are enriched for older males.'),
    ('2. Clone count (clonal diversity): ',
     'Recovered and COVID Naive individuals have 5-8x more unique clones than Mild or Severe cases. '
     'This likely reflects both biological differences and technical variation across studies.'),
    ('3. Clone size (clonal expansion): ',
     'COVID Naive subjects show the largest average clone sizes (median 114 sequences/clone), '
     '3x higher than Mild cases (36).'),
    ('4. Repertoire dominance (Top X fractions): ',
     'Moderate disease shows the most oligoclonal repertoire (Top 10 = 22.6%), while Mild cases '
     'are the most polyclonal (10.5%).'),
    ('5. CDR3 length: ',
     'Severe cases have the longest CDR3 in their top 10 clones (median 17.6 AA), while COVID Naive '
     'show the shortest (15.0 AA). Differences converge at the Top 1000 level.'),
    ('6. Sex differences: ',
     'Minimal differences in clone count, size, or CDR3 length between males and females.'),
    ('7. Age effects: ',
     'No strong age-dependent trends observed, though age x disease confounding limits interpretation.'),
    ('8. Caveats: ',
     'Disease category confounds with study of origin. Differences in sequencing depth, library '
     'preparation, and sample timing may contribute to variation. Results are exploratory.'),
]

for bold, rest in findings:
    p = doc.add_paragraph()
    run = p.add_run(bold)
    run.bold = True
    p.add_run(rest)
    p.paragraph_format.space_after = Pt(8)

doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run(
    'Generated from iReceptor Statistics API.\n'
    'Analysis scripts: metadata_visualization.R, clonal_analysis.R, cdr3_analysis.R'
)
run.italic = True
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

output_path = 'Immune_Repertoire_Results_v2.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
