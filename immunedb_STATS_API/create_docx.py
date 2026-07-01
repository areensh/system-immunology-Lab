from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# --- Helper functions ---
def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    return h

def add_figure(img_path, caption, width=Inches(5.5)):
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=width)
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True
        cap.runs[0].font.size = Pt(9)
        cap.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    else:
        doc.add_paragraph(f"[Image not found: {img_path}]")

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Shading Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()

def add_body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_bold_body(bold_part, rest):
    p = doc.add_paragraph()
    run = p.add_run(bold_part)
    run.bold = True
    p.add_run(rest)
    p.paragraph_format.space_after = Pt(6)

plots = "plots"
clonal = "plots/clonal"

# ============================================================
# TITLE PAGE
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading('Immune Repertoire Analysis', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_heading('Results Document', level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in subtitle.runs:
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('\nDate: ').bold = True
info.add_run('June 14, 2026\n')
info.add_run('Data Source: ').bold = True
info.add_run('iReceptor AIRR-seq Data Commons (COVID-19 studies)\n')
info.add_run('Total Repertoires: ').bold = True
info.add_run('100 (85 with complete metadata)')

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS (manual)
# ============================================================
add_heading('Table of Contents', level=1)
toc_items = [
    "1. Cohort Description & Metadata Overview",
    "   1.1 Data Source & Study Composition",
    "   1.2 Disease Category Distribution",
    "   1.3 Sex Distribution",
    "   1.4 Age Distribution",
    "   1.5 Cross-Stratification: Disease x Sex",
    "   1.6 Cross-Stratification: Disease x Age",
    "   1.7 Study x Disease Composition",
    "   1.8 Heatmap: Disease x Age x Sex",
    "2. Clonal Analysis",
    "   2.1 Clone Count",
    "   2.2 Clone Size",
    "   2.3 Top X Clone Fraction (Repertoire Dominance)",
    "   2.4 CDR3 Length (Amino Acid)",
    "3. Summary of Key Findings",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)

doc.add_page_break()

# ============================================================
# SECTION 1: COHORT DESCRIPTION
# ============================================================
add_heading('1. Cohort Description & Metadata Overview', level=1)

# 1.1
add_heading('1.1 Data Source & Study Composition', level=2)
add_body(
    'The dataset consists of 100 repertoires drawn from 3 COVID-19 studies '
    'retrieved via the iReceptor AIRR-seq Data Commons API. Of these, 85 repertoires '
    'had complete metadata (age, sex, disease stage) and were used for all downstream '
    'analyses. The remaining 15 lacked metadata fields.'
)
add_figure(f'{plots}/01_subjects_per_dataset.png',
           'Figure 1.1. Number of subjects per dataset (study).')

# 1.2
add_heading('1.2 Disease Category Distribution', level=2)
add_body(
    'Raw disease stage annotations were harmonized into 5 categories:'
)
add_table(
    ['Category', 'Description', 'N'],
    [
        ['COVID Naive', 'No prior SARS-CoV-2 infection', '8'],
        ['Mild', 'Non-severe / mild COVID-19', '38'],
        ['Moderate', 'Stable or improving COVID-19', '9'],
        ['Recovered', 'Post-COVID recovery', '12'],
        ['Severe', 'Severe / hypoxemic COVID-19', '18'],
    ]
)
add_body(
    'The Mild group is the largest (n=38, 45%), followed by Severe (n=18, 21%), '
    'Recovered (n=12, 14%), Moderate (n=9, 11%), and COVID Naive (n=8, 9%).'
)
add_figure(f'{plots}/02a_subjects_per_disease_category.png',
           'Figure 1.2. Number of subjects per disease category.')

# 1.3
add_heading('1.3 Sex Distribution', level=2)
add_body('The cohort has a slight male skew: 51 males (60%) and 34 females (40%).')
add_figure(f'{plots}/02b_subjects_per_sex.png',
           'Figure 1.3. Number of subjects by sex.')

# 1.4
add_heading('1.4 Age Distribution', level=2)
add_body('Age ranges from 18 to 88 years (median ~ 55). The age group breakdown is:')
add_table(
    ['Age Group', 'N'],
    [['18-30', '10'], ['31-50', '25'], ['51-65', '29'], ['66+', '21']]
)
add_body('The cohort skews older, with 59% of participants aged 51+.')
add_figure(f'{plots}/02c_age_distribution.png',
           'Figure 1.4a. Age distribution (histogram).')
add_figure(f'{plots}/02d_subjects_per_age_group.png',
           'Figure 1.4b. Number of subjects per age group.')

# 1.5
add_heading('1.5 Cross-Stratification: Disease x Sex', level=2)
add_body(
    'Males are overrepresented in Severe (12M vs 6F) and Mild (22M vs 16F) groups. '
    'COVID Naive is balanced (4M, 4F). Recovered has more males (8M vs 4F).'
)
add_figure(f'{plots}/03a_disease_by_sex.png',
           'Figure 1.5. Disease category distribution stratified by sex.')

# 1.6
add_heading('1.6 Cross-Stratification: Disease x Age', level=2)
add_body(
    'Older participants (51-65, 66+) are concentrated in the Severe group. '
    'The Mild group has the most even age distribution. COVID Naive participants tend to be younger.'
)
add_figure(f'{plots}/03b_disease_by_age_group.png',
           'Figure 1.6. Disease category distribution stratified by age group.')

# 1.7
add_heading('1.7 Study x Disease Composition', level=2)
add_body(
    'Different studies contributed different disease categories - disease confounds with '
    'study of origin, which is an important consideration for interpretation.'
)
add_figure(f'{plots}/03c_study_by_disease.png',
           'Figure 1.7. Study-level composition by disease category.')

# 1.8
add_heading('1.8 Heatmap: Disease x Age x Sex', level=2)
add_figure(f'{plots}/03d_age_by_disease_boxplot.png',
           'Figure 1.8a. Age distribution by disease category (boxplot).')
add_figure(f'{plots}/03e_heatmap_disease_age_sex.png',
           'Figure 1.8b. Heatmap of participant counts across disease category, age group, and sex.')

doc.add_page_break()

# ============================================================
# SECTION 2: CLONAL ANALYSIS
# ============================================================
add_heading('2. Clonal Analysis', level=1)

# 2.1 Clone Count
add_heading('2.1 Clone Count', level=2)
add_body(
    'Clone count represents the number of unique B-cell clones identified in each repertoire. '
    'Higher clone counts reflect greater clonal diversity.'
)

add_heading('2.1.1 Clone Count by Disease Category', level=3)
add_table(
    ['Disease Category', 'N', 'Median', 'Mean', 'Min', 'Max'],
    [
        ['COVID Naive', '8', '16,682', '18,086', '8,156', '29,178'],
        ['Mild', '38', '2,353', '2,593', '95', '5,997'],
        ['Moderate', '9', '5,085', '9,327', '3,459', '22,480'],
        ['Recovered', '12', '21,527', '19,661', '5,713', '33,699'],
        ['Severe', '18', '3,423', '6,782', '1,319', '22,452'],
    ]
)
add_bold_body('Key finding: ',
    'Recovered and COVID Naive individuals have the highest clone counts (median ~17K-22K), '
    'suggesting greater clonal diversity. Mild cases have the lowest (median ~2,350), likely '
    'reflecting differences in sequencing depth or sampling across studies rather than a '
    'biological signal alone.'
)
add_figure(f'{clonal}/cc_by_disease.png',
           'Figure 2.1.1. Clone count by disease category.')

add_heading('2.1.2 Clone Count by Sex', level=3)
add_table(
    ['Sex', 'N', 'Median', 'Mean'],
    [['Female', '34', '4,002', '7,073'], ['Male', '51', '4,244', '8,719']]
)
add_body('Clone counts are comparable between sexes, with males showing a slightly higher mean driven by outliers.')
add_figure(f'{clonal}/cc_by_sex.png',
           'Figure 2.1.2. Clone count by sex.')

add_heading('2.1.3 Clone Count by Age Group', level=3)
add_figure(f'{clonal}/cc_by_age.png',
           'Figure 2.1.3. Clone count by age group.')

add_heading('2.1.4 Clone Count by Disease x Sex', level=3)
add_figure(f'{clonal}/cc_by_disease_sex.png',
           'Figure 2.1.4. Clone count by disease category and sex.')

add_heading('2.1.5 Clone Count by Disease x Age', level=3)
add_figure(f'{clonal}/cc_by_disease_age.png',
           'Figure 2.1.5. Clone count by disease category and age group.')

doc.add_page_break()

# 2.2 Clone Size
add_heading('2.2 Clone Size', level=2)
add_body(
    'Clone size represents the mean number of sequences per clone within each repertoire. '
    'Larger clone sizes indicate greater clonal expansion (i.e., individual clones have proliferated more).'
)

add_heading('2.2.1 Mean Clone Size by Disease Category', level=3)
add_table(
    ['Disease Category', 'N', 'Median (Mean CS)', 'Mean (Mean CS)'],
    [
        ['COVID Naive', '8', '114.0', '126.0'],
        ['Mild', '38', '36.4', '48.5'],
        ['Moderate', '9', '63.5', '60.0'],
        ['Recovered', '12', '66.9', '79.3'],
        ['Severe', '18', '53.4', '52.0'],
    ]
)
add_bold_body('Key finding: ',
    'COVID Naive individuals show the highest mean clone size (median 114), suggesting that even '
    'without COVID infection, their baseline repertoire features larger expanded clones. Mild cases '
    'show the smallest clones (median 36.4), consistent with a less intense immune response or '
    'earlier sampling. Severe and Moderate groups show intermediate expansion.'
)
add_figure(f'{clonal}/cs_by_disease.png',
           'Figure 2.2.1. Mean clone size by disease category.')

add_heading('2.2.2 Mean Clone Size by Sex', level=3)
add_figure(f'{clonal}/cs_by_sex.png',
           'Figure 2.2.2. Mean clone size by sex.')

add_heading('2.2.3 Mean Clone Size by Age Group', level=3)
add_figure(f'{clonal}/cs_by_age.png',
           'Figure 2.2.3. Mean clone size by age group.')

add_heading('2.2.4 Mean Clone Size by Disease x Sex', level=3)
add_figure(f'{clonal}/cs_by_disease_sex.png',
           'Figure 2.2.4. Mean clone size by disease category and sex.')

doc.add_page_break()

# 2.3 Top X Clone Fraction
add_heading('2.3 Top X Clone Fraction (Repertoire Dominance)', level=2)
add_body(
    'The Top X clone fraction measures the proportion of total sequence copies accounted for by '
    'the top 10, 100, or 1,000 clones. Higher fractions indicate greater oligoclonal dominance - '
    'a few clones dominate the repertoire.'
)

add_heading('2.3.1 Top X Fractions by Disease Category', level=3)
add_table(
    ['Disease Category', 'N', 'Top 10 (median)', 'Top 100 (median)', 'Top 1000 (median)'],
    [
        ['COVID Naive', '8', '14.4%', '61.0%', '83.3%'],
        ['Mild', '38', '10.5%', '30.4%', '72.7%'],
        ['Moderate', '9', '22.6%', '45.5%', '74.1%'],
        ['Recovered', '12', '16.2%', '40.3%', '76.9%'],
        ['Severe', '18', '16.6%', '35.6%', '74.4%'],
    ]
)

add_bold_body('Key findings:', '')
bullets = [
    'Moderate disease shows the highest Top 10 dominance (22.6%), suggesting highly focused clonal responses in patients with stable/improving disease.',
    'COVID Naive has the highest Top 100 and Top 1000 fractions (61% and 83%), indicating an oligoclonal baseline repertoire.',
    'Mild cases show the lowest dominance across all tiers (10.5% Top 10), consistent with a more polyclonal, less focused response.',
]
for b in bullets:
    p = doc.add_paragraph(b, style='List Bullet')
    p.paragraph_format.space_after = Pt(4)

add_figure(f'{clonal}/topX_by_disease.png',
           'Figure 2.3.1. Top X clone fractions by disease category (faceted by Top 10, 100, 1000).',
           width=Inches(6))

add_heading('2.3.2 Top X Fractions by Sex', level=3)
add_figure(f'{clonal}/topX_by_sex.png',
           'Figure 2.3.2. Top X clone fractions by sex.', width=Inches(5.5))

add_heading('2.3.3 Top X Fractions by Age Group', level=3)
add_figure(f'{clonal}/topX_by_age.png',
           'Figure 2.3.3. Top X clone fractions by age group.', width=Inches(5.5))

add_heading('2.3.4 Top 10 Clone Fraction by Disease x Sex', level=3)
add_figure(f'{clonal}/top10_by_disease_sex.png',
           'Figure 2.3.4. Top 10 clone fraction by disease category and sex.')

add_heading('2.3.5 Top 10 Clone Fraction by Disease x Age', level=3)
add_figure(f'{clonal}/top10_by_disease_age.png',
           'Figure 2.3.5. Top 10 clone fraction by disease category and age group.')

doc.add_page_break()

# ============================================================
# SECTION 2.4: CDR3 LENGTH
# ============================================================
cdr3 = "plots/cdr3"

add_heading('2.4 CDR3 Length (Amino Acid)', level=2)
add_body(
    'CDR3 (Complementarity-Determining Region 3) length is a key structural feature of antibodies '
    'that determines antigen-binding specificity. Longer CDR3 regions can access recessed epitopes '
    'and are often associated with broadly neutralizing antibodies. Here we analyze the average '
    'CDR3 length (in amino acids) of the top 10, 100, and 1,000 most expanded clones per repertoire.'
)

add_heading('2.4.1 CDR3 Length by Disease Category', level=3)
add_table(
    ['Disease Category', 'N', 'Top 10 (median)', 'Top 100 (median)', 'Top 1000 (median)'],
    [
        ['COVID Naive', '8', '15.0 AA', '16.8 AA', '17.4 AA'],
        ['Mild', '38', '16.5 AA', '16.9 AA', '17.0 AA'],
        ['Moderate', '9', '17.3 AA', '17.2 AA', '17.5 AA'],
        ['Recovered', '12', '15.8 AA', '17.2 AA', '17.5 AA'],
        ['Severe', '18', '17.6 AA', '17.4 AA', '17.2 AA'],
    ]
)
add_bold_body('Key finding: ',
    'Severe cases show the longest CDR3 in their top 10 clones (median 17.6 AA), while COVID Naive '
    'show the shortest (median 15.0 AA). However, CDR3 lengths converge across disease categories '
    'when considering the top 1000 clones (~17.0-17.5 AA), suggesting that the CDR3 length '
    'differences are primarily driven by the most dominant clones.'
)
add_figure(f'{cdr3}/cdr3_by_disease.png',
           'Figure 2.4.1. Average CDR3 length (AA) by disease category (faceted by clone tier).',
           width=Inches(6))

add_heading('2.4.2 CDR3 Length by Sex', level=3)
add_body(
    'CDR3 lengths are comparable between males (median 17.0 AA) and females (median 17.1 AA) '
    'for the top 10 clones, with no meaningful sex-based differences observed.'
)
add_figure(f'{cdr3}/cdr3_by_sex.png',
           'Figure 2.4.2. Average CDR3 length (AA) by sex.', width=Inches(5.5))

add_heading('2.4.3 CDR3 Length by Age Group', level=3)
add_figure(f'{cdr3}/cdr3_by_age.png',
           'Figure 2.4.3. Average CDR3 length (AA) by age group.', width=Inches(5.5))

add_heading('2.4.4 CDR3 Length (Top 10) by Disease x Sex', level=3)
add_figure(f'{cdr3}/cdr3_top10_by_disease_sex.png',
           'Figure 2.4.4. Average CDR3 length of top 10 clones by disease category and sex.')

add_heading('2.4.5 CDR3 Length (Top 10) by Disease x Age', level=3)
add_figure(f'{cdr3}/cdr3_top10_by_disease_age.png',
           'Figure 2.4.5. Average CDR3 length of top 10 clones by disease category and age group.')

add_heading('2.4.6 CDR3 Length Across Clone Tiers', level=3)
add_body(
    'This figure compares CDR3 lengths across the three clone tiers (Top 10, 100, 1000) '
    'simultaneously, showing how the most dominant clones differ from the broader repertoire.'
)
add_figure(f'{cdr3}/cdr3_tiers_by_disease.png',
           'Figure 2.4.6. CDR3 length across clone tiers by disease category.')

doc.add_page_break()

# ============================================================
# SECTION 3: SUMMARY
# ============================================================
add_heading('3. Summary of Key Findings', level=1)

findings = [
    ('1. Cohort composition: ',
     '85 participants across 5 disease categories (Mild is the largest group at 45%). '
     'The cohort skews male (60%) and older (59% aged 51+). Severe cases are enriched for older males.'),
    ('2. Clone count (clonal diversity): ',
     'Recovered and COVID Naive individuals have 5-8x more unique clones than Mild or Severe cases. '
     'This likely reflects both biological differences (post-infection expansion, baseline diversity) '
     'and technical variation (sequencing depth across studies).'),
    ('3. Clone size (clonal expansion): ',
     'COVID Naive subjects show the largest average clone sizes (median 114 sequences/clone), '
     '3x higher than Mild cases (36). This suggests substantial baseline clonal expansion even '
     'without COVID infection.'),
    ('4. Repertoire dominance (Top X fractions): ',
     'Moderate disease shows the most oligoclonal repertoire (Top 10 clones = 22.6% of copies), '
     'while Mild cases are the most polyclonal (10.5%). This pattern is consistent with focused '
     'immune responses in more symptomatic disease.'),
    ('5. CDR3 length: ',
     'Severe cases have the longest CDR3 in their top 10 clones (median 17.6 AA), while COVID Naive '
     'show the shortest (15.0 AA). CDR3 lengths converge across groups at the Top 1000 level (~17 AA), '
     'indicating differences are driven by the most dominant expanded clones.'),
    ('6. Sex differences: ',
     'Minimal differences in clone count, size, or CDR3 length between males and females across all disease categories.'),
    ('7. Age effects: ',
     'No strong age-dependent trends in clonal metrics were observed, though the age x disease '
     'confound (older patients in Severe group) limits interpretation.'),
    ('8. Caveats: ',
     'Disease category confounds with study of origin. Differences in sequencing depth, library '
     'preparation, and sample timing across studies may contribute to observed variation. These '
     'results should be interpreted as exploratory.'),
]

for bold, rest in findings:
    p = doc.add_paragraph()
    run = p.add_run(bold)
    run.bold = True
    p.add_run(rest)
    p.paragraph_format.space_after = Pt(8)

doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run(
    'Generated from iReceptor AIRR Data Commons API data.\n'
    'Analysis scripts: metadata_visualization.R, clonal_analysis.R, cdr3_analysis.R'
)
run.italic = True
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

output_path = 'Immune_Repertoire_Analysis_Results.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
