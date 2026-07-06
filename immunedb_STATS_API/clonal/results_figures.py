from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
import os

BASE = "/home/user/system-immunology-Lab/immunedb_STATS_API"
os.chdir(BASE + "/clonal")

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

for s in ['Heading 1', 'Heading 2', 'Heading 3']:
    doc.styles[s].font.name = 'Times New Roman'
    doc.styles[s].font.color.rgb = RGBColor(0, 0, 0)

def add_landscape_section(doc):
    new_section = doc.add_section(start_type=2)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = Inches(11)
    new_section.page_height = Inches(8.5)
    new_section.left_margin = Inches(0.5)
    new_section.right_margin = Inches(0.5)

def add_portrait_section(doc):
    new_section = doc.add_section(start_type=2)
    new_section.orientation = WD_ORIENT.PORTRAIT
    new_section.page_width = Inches(8.5)
    new_section.page_height = Inches(11)
    new_section.left_margin = Inches(1)
    new_section.right_margin = Inches(1)

def add_figure(doc, path, caption, width=7.0, landscape=False):
    if os.path.exists(path):
        if landscape:
            add_landscape_section(doc)
            w = 10.0
        else:
            w = width
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(w))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.font.size = Pt(10)
        run.font.italic = True
        doc.add_paragraph()
        if landscape:
            add_portrait_section(doc)

def add_table(doc, headers, rows, caption=None):
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.font.size = Pt(10)
        run.font.bold = True

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Shading Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(9)

    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()

# ============================================================
# TITLE
# ============================================================
title = doc.add_heading('Results', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

intro = doc.add_paragraph()
intro.add_run(
    'Clonal analysis was performed on immune repertoire data from 103 subjects across '
    '7 studies (CD1, CD2, CD3, CVX1, CVX2, HC1, GT1), using peripheral blood samples only. '
    'Subjects were categorized into six disease groups: Severe (n=27), Mild (n=41), '
    'Moderate (n=9), Recovered (n=12), COVID Naive (n=8), and Healthy (n=6). '
    'Disease labels from individual studies were harmonized into these categories. '
    'Statistics were computed using the iReceptor Statistics API (v0.3.0) with a metadata '
    'fingerprint grouping approach that correctly aggregates per-tissue data across samples.'
)

META = BASE + "/metadata/plots"

# ============================================================
# 0. METADATA / COHORT DESCRIPTION
# ============================================================
doc.add_heading('1. Cohort Metadata and Demographics', level=2)

doc.add_paragraph(
    'The study cohort was assembled from seven independent immune repertoire datasets '
    'retrieved via the iReceptor Gateway: CD1 (Covid19_db3, n=51), CD2 (covid_db2, n=19), '
    'CD3 (covid19, n=13), CVX1 (vaccine2, n=12), CVX2 (covid_vaccine_new, n=5), '
    'HC1 (lp16_Igblast), and GT1 (sykesIgblast2020). After excluding control samples '
    '(covid_vaccine_new-Fb, Water) and non-human subjects (lp16_Igblast-D159, D154, Hu-1), '
    'and removing healthy subjects from the CD3 study (H3, H4, H8), the final analysis '
    'included 103 subjects with peripheral blood data.'
)

doc.add_heading('1.1 Metadata Availability per Dataset', level=3)
doc.add_paragraph(
    'Figure 1 summarizes the metadata fields available across datasets. All datasets '
    'provided disease stage and tissue information. Sex metadata was available for all '
    'datasets except CD3 and HC1. Age data was available for CD1, CD2, CVX1, and CVX2.'
)
add_figure(doc, META + '/00_metadata_per_dataset.png',
           'Figure 1. Metadata availability per dataset. Bar heights indicate the number '
           'of subjects with each metadata type available.')

doc.add_heading('1.2 Subjects per Dataset', level=3)
doc.add_paragraph(
    'Figure 2 shows the number of subjects contributed by each dataset. CD1 was the '
    'largest contributor (51 subjects), followed by CD2 (19) and CD3 (13).'
)
add_figure(doc, META + '/01_subjects_per_dataset.png',
           'Figure 2. Number of subjects per dataset.')

doc.add_heading('1.3 Disease Category Distribution', level=3)
doc.add_paragraph(
    'Disease labels from individual studies were harmonized into six categories: '
    'Severe (n=27), Mild (n=41), Moderate (n=9), Recovered (n=12), COVID Naive (n=8), '
    'and Healthy (n=6). The Mild category was the most represented, combining subjects '
    'labeled as "mild" (CD1) and "non-severe" (CD3). The Severe category included '
    'subjects with "severe" (CD1, CD3) and "Early phase hypoxaemia" (CD2) labels. '
    'The Moderate category comprised CD2 subjects with "Early phase-Stable" or '
    '"Early phase-Improving" labels.'
)
add_figure(doc, META + '/02a_subjects_per_disease_category.png',
           'Figure 3. Number of subjects per harmonized disease category.')

doc.add_heading('1.4 Sex Distribution', level=3)
doc.add_paragraph(
    'Among subjects with available sex metadata (n=86), 52 were male and 34 were female. '
    'Fourteen subjects (CD3 and HC1 studies) lacked sex annotation.'
)
add_figure(doc, META + '/02b_subjects_per_sex.png',
           'Figure 4. Sex distribution across the cohort.')

doc.add_heading('1.5 Age Distribution', level=3)
doc.add_paragraph(
    'Age data was available for 86 subjects (median 48 years, range 18-87). '
    'The age distribution was roughly uniform across the 31-50 and 51-65 groups, '
    'with fewer subjects in the youngest (18-30) and oldest (66+) brackets.'
)
add_figure(doc, META + '/02c_age_distribution.png',
           'Figure 5. Age distribution histogram. Red dashed line indicates the median age.')
add_figure(doc, META + '/02d_subjects_per_age_group.png',
           'Figure 6. Number of subjects per age group.')

doc.add_heading('1.6 Cross-tabulations', level=3)
doc.add_paragraph(
    'Figures 7-10 show the relationships between disease category, sex, age, and study. '
    'The sex ratio was relatively balanced within each disease category. Age distributions '
    'varied across disease groups, with Moderate and Recovered subjects tending to be older.'
)
add_figure(doc, META + '/03a_disease_by_sex.png',
           'Figure 7. Disease category by sex.')
add_figure(doc, META + '/03b_disease_by_age_group.png',
           'Figure 8. Disease category by age group.')
add_figure(doc, META + '/03c_study_by_disease.png',
           'Figure 9. Study composition by disease category.')
add_figure(doc, META + '/03d_age_by_disease_boxplot.png',
           'Figure 10. Age distribution by disease category (box plots with individual data points).')
add_figure(doc, META + '/03e_heatmap_disease_age_sex.png',
           'Figure 11. Heatmap of subject counts by disease category, age group, and sex.')

# ============================================================
# 2. CLONE SIZE
# ============================================================
doc.add_page_break()
doc.add_heading('2. Clone Size Analysis', level=2)

doc.add_paragraph(
    'Clone size, defined as the number of sequence copies per clone, reflects the degree '
    'of clonal expansion in each subject\'s repertoire. We analyzed both the number of '
    'expanded clones (clones with size > 20) and the median clone size per subject.'
)

doc.add_heading('2.1 Number of Expanded Clones by Disease Category', level=3)
doc.add_paragraph(
    'Figure 12 shows the number of expanded clones (size > 20) across disease categories. '
    'COVID Naive and Healthy subjects showed higher counts of expanded clones compared to '
    'acute disease groups, suggesting pre-existing clonal expansions unrelated to SARS-CoV-2 infection.'
)
add_figure(doc, 'clone_size/plots/01_n_expanded_clones_by_disease.png',
           'Figure 12. Number of expanded clones (size > 20) by disease category. '
           'Box plots show median and IQR; red diamond indicates mean with SD error bars. '
           'Points are colored by original disease label.')

doc.add_heading('2.2 Median Clone Size by Disease Category', level=3)
doc.add_paragraph(
    'Figure 13 presents the median clone size per subject across disease categories. '
    'The median clone size was relatively uniform across groups, with COVID Naive subjects '
    'showing slightly elevated values.'
)
add_figure(doc, 'clone_size/plots/02_median_clone_size_by_disease.png',
           'Figure 13. Median clone size per subject by disease category.')

doc.add_heading('2.3 Clone Size by Age and Disease', level=3)
doc.add_paragraph(
    'Figures 14 and 15 examine the number of expanded clones and median clone size '
    'stratified by age group within each disease category.'
)
add_figure(doc, 'clone_size/plots/03_clone_count_by_age_disease.png',
           'Figure 14. Number of expanded clones by age group, faceted by disease category.')
add_figure(doc, 'clone_size/plots/05_clone_size_by_age_disease.png',
           'Figure 15. Median clone size by age group, faceted by disease category.')

doc.add_heading('2.4 Clone Size by Sex and Disease', level=3)
doc.add_paragraph(
    'Figures 16 and 17 examine sex-based differences in clonal expansion within each '
    'disease category.'
)
add_figure(doc, 'clone_size/plots/04_clone_count_by_sex_disease.png',
           'Figure 16. Number of expanded clones by sex, faceted by disease category.')
add_figure(doc, 'clone_size/plots/06_clone_size_by_sex_disease.png',
           'Figure 17. Median clone size by sex, faceted by disease category.')

# ============================================================
# 2. CLONE COUNT (CLONAL DIVERSITY)
# ============================================================
doc.add_page_break()
doc.add_heading('3. Clonal Diversity (Clone Count)', level=2)

doc.add_paragraph(
    'Clone count, representing the number of unique clones per subject, serves as a measure '
    'of clonal diversity. Higher clone counts indicate greater repertoire diversity.'
)

doc.add_heading('3.1 Clone Count by Disease Category', level=3)
doc.add_paragraph(
    'Figure 18 shows the distribution of unique clone counts across disease categories. '
    'Recovered and COVID Naive subjects exhibited markedly higher clone counts '
    '(median 21,526 and 16,682, respectively) compared to Mild (2,478) and Severe (5,089) '
    'groups, suggesting repertoire expansion following infection or vaccination. '
    'The Healthy group showed high variability (range 650-61,136).'
)
add_figure(doc, 'clone_count/plots/07_clone_count_by_disease.png',
           'Figure 18. Clone count (number of unique clones) by disease category. '
           'Box plots show median and IQR; red diamond indicates mean with SD error bars.')

add_table(doc,
    ['Disease Category', 'n', 'Median', 'Mean', 'Min', 'Max'],
    [
        ['Severe', '27', '5,089', '9,529', '1,319', '35,137'],
        ['Mild', '41', '2,478', '3,127', '95', '22,565'],
        ['Moderate', '9', '5,085', '9,327', '3,459', '22,480'],
        ['Recovered', '12', '21,526', '19,661', '5,713', '33,699'],
        ['COVID Naive', '8', '16,682', '18,086', '8,156', '29,178'],
        ['Healthy', '6', '1,707', '14,443', '650', '61,136'],
    ],
    caption='Table 1. Clone count summary statistics by disease category.'
)

doc.add_heading('3.2 Clone Count by Age Group and Disease', level=3)
doc.add_paragraph(
    'Figure 19 displays clone count distributions stratified by age group within each '
    'disease category.'
)
add_figure(doc, 'clone_count/plots/08_clone_count_by_age_disease.png',
           'Figure 19. Clone count by age group, faceted by disease category.')

doc.add_heading('3.3 Clone Count by Sex and Disease', level=3)
doc.add_paragraph(
    'Figure 20 shows clone count distributions by sex within each disease category. '
    'No substantial sex-based differences were observed across disease groups.'
)
add_figure(doc, 'clone_count/plots/09_clone_count_by_sex_disease.png',
           'Figure 20. Clone count by sex, faceted by disease category.')

add_table(doc,
    ['Disease Category', 'Sex', 'n', 'Median Clone Count'],
    [
        ['Severe', 'Female', '5', '3,552'],
        ['Severe', 'Male', '14', '3,567'],
        ['Mild', 'Female', '19', '2,140'],
        ['Mild', 'Male', '19', '2,478'],
        ['Moderate', 'Female', '3', '4,930'],
        ['Moderate', 'Male', '6', '5,102'],
        ['Recovered', 'Female', '3', '23,711'],
        ['Recovered', 'Male', '9', '20,876'],
        ['COVID Naive', 'Female', '4', '18,745'],
        ['COVID Naive', 'Male', '4', '16,364'],
        ['Healthy', 'Female', '1', '1,972'],
        ['Healthy', 'Male', '5', '1,442'],
    ],
    caption='Table 2. Clone count by disease category and sex.'
)

# ============================================================
# 3. TOP-X CLONE PROPORTION
# ============================================================
doc.add_page_break()
doc.add_heading('4. Top-X Clone Proportion Analysis', level=2)

doc.add_paragraph(
    'The proportion of total repertoire copies held by the top expanded clones '
    '(Top 10, Top 100, and Top 1000) provides a measure of clonal dominance. '
    'Higher proportions indicate that a small number of clones dominate the repertoire.'
)

doc.add_heading('4.1 Top-X Proportion by Disease Category', level=3)
doc.add_paragraph(
    'Figure 21 shows the proportion of total copies captured by the top 10, 100, and 1000 '
    'clones across disease categories. Healthy subjects showed notably high Top 1000 '
    'concentration (median 99.3%), indicating that nearly all copies belong to the top 1000 '
    'clones, consistent with a repertoire dominated by relatively few expanded clones. '
    'Moderate disease subjects showed the highest Top 10 dominance (median 23.1%).'
)
add_figure(doc, 'topX/plots/10_topX_stacked_by_disease.png',
           'Figure 21. Clonal dominance: fraction of total repertoire copies held by '
           'Top 10 (red), Top 11-100 (orange), Top 101-1000 (green), and remaining (blue) '
           'clones per subject, faceted by disease category and sorted by Top 10 dominance.',
           landscape=True)

add_table(doc,
    ['Disease Category', 'n', 'Top 10 (%)', 'Top 100 (%)', 'Top 1000 (%)'],
    [
        ['Severe', '27', '15.8', '34.4', '70.8'],
        ['Mild', '41', '10.6', '33.3', '75.2'],
        ['Moderate', '9', '23.1', '46.4', '74.5'],
        ['Recovered', '12', '13.7', '35.9', '62.0'],
        ['COVID Naive', '8', '9.6', '36.0', '74.0'],
        ['Healthy', '6', '16.9', '58.7', '99.3'],
    ],
    caption='Table 3. Median proportion of total copies held by top-X clones.'
)

doc.add_heading('4.2 Top-X Proportion by Age and Disease', level=3)
doc.add_paragraph(
    'Figure 22 presents the top-X clone proportions stratified by age group within each '
    'disease category.'
)
add_figure(doc, 'topX/plots/11_topX_stacked_by_age_disease.png',
           'Figure 22. Clonal dominance stacked bars per subject, faceted by disease '
           'category and age group, sorted by Top 10 dominance.',
           landscape=True)

doc.add_heading('4.3 Top-X Proportion by Sex and Disease', level=3)
doc.add_paragraph(
    'Figure 23 shows top-X clone proportions by sex within each disease category.'
)
add_figure(doc, 'topX/plots/11b_topX_stacked_by_sex_disease.png',
           'Figure 23. Clonal dominance stacked bars per subject, faceted by disease '
           'category and sex, sorted by Top 10 dominance.',
           landscape=True)

add_table(doc,
    ['Disease Category', 'Sex', 'n', 'Top 10 (%)', 'Top 100 (%)'],
    [
        ['Severe', 'Female', '5', '12.4', '38.6'],
        ['Severe', 'Male', '14', '19.2', '35.9'],
        ['Mild', 'Female', '19', '10.5', '33.4'],
        ['Mild', 'Male', '19', '10.9', '31.8'],
        ['Moderate', 'Female', '3', '23.1', '46.4'],
        ['Moderate', 'Male', '6', '25.2', '47.9'],
        ['Recovered', 'Female', '3', '8.6', '20.3'],
        ['Recovered', 'Male', '9', '15.0', '36.4'],
        ['COVID Naive', 'Female', '4', '11.0', '39.1'],
        ['COVID Naive', 'Male', '4', '7.4', '36.0'],
        ['Healthy', 'Female', '1', '18.8', '68.3'],
        ['Healthy', 'Male', '5', '15.1', '49.0'],
    ],
    caption='Table 4. Median top-X clone proportions by disease category and sex.'
)

# ============================================================
# 4. CDR3 AA LENGTH
# ============================================================
doc.add_page_break()
doc.add_heading('5. CDR3 Amino Acid Length Analysis', level=2)

doc.add_paragraph(
    'The CDR3 (Complementarity-Determining Region 3) amino acid length is a key structural '
    'feature of antibodies that influences antigen binding specificity. We analyzed the '
    'average CDR3 AA length of the top expanded clones (Top 10, Top 100, and Top 1000) '
    'to examine whether disease severity is associated with differences in CDR3 length '
    'of dominant clones.'
)

doc.add_heading('5.1 CDR3 Length by Disease Category', level=3)
doc.add_paragraph(
    'Figure 24 shows the average CDR3 AA length of the top expanded clones across disease '
    'categories. Severe cases exhibited the longest CDR3 lengths in Top 10 clones '
    '(median 17.7 AA), while Healthy subjects showed the shortest (median 13.8 AA). '
    'This difference was most pronounced in the Top 10 tier and diminished progressively '
    'with larger clone sets (Top 100, Top 1000), suggesting that the most expanded clones '
    'in severe disease tend to use longer CDR3 regions.'
)
add_figure(doc, 'cdr3/plots/12_cdr3_length_by_disease.png',
           'Figure 24. Average CDR3 amino acid length per subject for Top 10 (red), '
           'Top 100 (orange), and Top 1000 (green) expanded clones, faceted by disease '
           'category (columns) and clone tier (rows). Subjects sorted by Top 10 CDR3 length.',
           landscape=True)

add_table(doc,
    ['Disease Category', 'n', 'Top 10 (AA)', 'Top 100 (AA)', 'Top 1000 (AA)'],
    [
        ['Severe', '27', '17.7', '17.0', '17.1'],
        ['Mild', '41', '16.3', '16.9', '16.9'],
        ['Moderate', '9', '17.3', '17.2', '17.4'],
        ['Recovered', '12', '16.3', '16.6', '17.0'],
        ['COVID Naive', '8', '15.8', '16.6', '16.9'],
        ['Healthy', '6', '13.8', '14.6', '16.1'],
    ],
    caption='Table 5. Median average CDR3 AA length of top-X clones by disease category.'
)

doc.add_heading('5.2 CDR3 Length by Age and Disease', level=3)
doc.add_paragraph(
    'Figure 25 shows CDR3 AA length distributions stratified by age group within each '
    'disease category.'
)
add_figure(doc, 'cdr3/plots/13_cdr3_length_by_age_disease.png',
           'Figure 25. Average CDR3 AA length per subject, faceted by disease category, '
           'age group (columns), and clone tier (rows).',
           landscape=True)

doc.add_heading('5.3 CDR3 Length by Sex and Disease', level=3)
doc.add_paragraph(
    'Figure 26 presents CDR3 AA length distributions by sex within each disease category.'
)
add_figure(doc, 'cdr3/plots/14_cdr3_length_by_sex_disease.png',
           'Figure 26. Average CDR3 AA length per subject, faceted by disease category, '
           'sex (columns), and clone tier (rows).',
           landscape=True)

add_table(doc,
    ['Disease Category', 'Sex', 'n', 'Top 10 (AA)', 'Top 100 (AA)'],
    [
        ['Severe', 'Female', '5', '17.5', '16.8'],
        ['Severe', 'Male', '14', '17.8', '17.5'],
        ['Mild', 'Female', '19', '16.3', '16.9'],
        ['Mild', 'Male', '19', '16.5', '16.9'],
        ['Moderate', 'Female', '3', '17.8', '17.4'],
        ['Moderate', 'Male', '6', '17.0', '17.2'],
        ['Recovered', 'Female', '3', '15.1', '16.5'],
        ['Recovered', 'Male', '9', '17.2', '16.7'],
        ['COVID Naive', 'Female', '4', '17.2', '16.7'],
        ['COVID Naive', 'Male', '4', '14.6', '16.5'],
        ['Healthy', 'Female', '1', '12.8', '14.0'],
        ['Healthy', 'Male', '5', '14.5', '14.9'],
    ],
    caption='Table 6. Median CDR3 AA length by disease category and sex.'
)

# ============================================================
# 6. SOMATIC HYPERMUTATION
# ============================================================
doc.add_page_break()
doc.add_heading('6. Somatic Hypermutation Analysis', level=2)

doc.add_paragraph(
    'Somatic hypermutation (SHM) is a key process in antibody affinity maturation, '
    'introducing point mutations in immunoglobulin variable regions. We analyzed the '
    'average number of mutations in the top expanded clones (Top 10, Top 100, and Top 1000) '
    'per subject to examine whether disease severity is associated with differences in '
    'mutation burden among dominant clones.'
)

doc.add_heading('6.1 Mutation Level by Disease Category', level=3)
doc.add_paragraph(
    'Figure 27 shows the average mutation count of the top expanded clones across disease '
    'categories. COVID Naive subjects exhibited the highest mutation levels in Top 10 clones '
    '(median 119.0), followed by Recovered (median 107.0), suggesting extensive affinity '
    'maturation in subjects with prior immune exposure. Mild subjects showed the lowest '
    'Top 10 mutation levels (median 48.6). The difference between tiers was most pronounced '
    'in COVID Naive subjects, where Top 10 clones carried substantially more mutations than '
    'Top 1000 clones, indicating strong selection of highly mutated dominant clones.'
)
add_figure(doc, 'mutation/plots/15_mutation_by_disease.png',
           'Figure 27. Average mutation count per subject for Top 10 (red), Top 100 (orange), '
           'and Top 1000 (green) expanded clones, faceted by disease category. '
           'Box plots show median and IQR; red diamond indicates mean with SD error bars.',
           landscape=True)

add_table(doc,
    ['Disease Category', 'n', 'Top 10', 'Top 100', 'Top 1000'],
    [
        ['Severe', '27', '80.2', '33.5', '16.9'],
        ['Mild', '41', '48.6', '25.5', '15.0'],
        ['Moderate', '9', '74.8', '30.3', '13.0'],
        ['Recovered', '12', '107.0', '60.1', '23.5'],
        ['COVID Naive', '8', '119.0', '92.9', '27.3'],
        ['Healthy', '6', '78.9', '35.5', '12.7'],
    ],
    caption='Table 7. Median average mutation count of top-X clones by disease category.'
)

doc.add_heading('6.2 Mutation Level by Age and Disease', level=3)
doc.add_paragraph(
    'Figure 28 shows mutation level distributions stratified by age group within each '
    'disease category.'
)
add_figure(doc, 'mutation/plots/16_mutation_by_age_disease.png',
           'Figure 28. Average mutation count per subject, faceted by disease category '
           '(columns) and age group (rows), for each clone tier.',
           landscape=True)

doc.add_heading('6.3 Mutation Level by Sex and Disease', level=3)
doc.add_paragraph(
    'Figure 29 presents mutation level distributions by sex within each disease category.'
)
add_figure(doc, 'mutation/plots/17_mutation_by_sex_disease.png',
           'Figure 29. Average mutation count per subject, faceted by disease category '
           '(columns) and sex (rows), for each clone tier.',
           landscape=True)

add_table(doc,
    ['Disease Category', 'Sex', 'n', 'Top 10', 'Top 100'],
    [
        ['Severe', 'Female', '5', '86.0', '41.4'],
        ['Severe', 'Male', '14', '82.0', '32.7'],
        ['Mild', 'Female', '19', '46.2', '25.5'],
        ['Mild', 'Male', '19', '48.6', '24.1'],
        ['Moderate', 'Female', '3', '113.0', '40.6'],
        ['Moderate', 'Male', '6', '71.3', '29.4'],
        ['Recovered', 'Female', '3', '101.0', '52.5'],
        ['Recovered', 'Male', '9', '120.0', '60.9'],
        ['COVID Naive', 'Female', '4', '115.0', '91.2'],
        ['COVID Naive', 'Male', '4', '121.0', '101.0'],
        ['Healthy', 'Female', '1', '82.1', '36.0'],
        ['Healthy', 'Male', '5', '75.8', '35.0'],
    ],
    caption='Table 8. Median mutation count by disease category and sex.'
)

# ============================================================
# SAVE
# ============================================================
outpath = '/home/user/system-immunology-Lab/immunedb_STATS_API/clonal/Clonal_Analysis_Results.docx'
doc.save(outpath)
print(f"Document saved to {outpath}")
